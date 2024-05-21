import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from docx import Document
import os

def save_to_docx(event=None):
    item_name = entry_item_name.get("1.0", tk.END).strip()
    calories = entry_calories.get("1.0", tk.END).strip()
    calories_from_fat = entry_calories_from_fat.get("1.0", tk.END).strip()
    fat = entry_fat.get("1.0", tk.END).strip()
    saturated_fat = entry_saturated_fat.get("1.0", tk.END).strip()
    trans_fat = entry_trans_fat.get("1.0", tk.END).strip()
    cholesterol = entry_cholesterol.get("1.0", tk.END).strip()
    sodium = entry_sodium.get("1.0", tk.END).strip()
    carbohydrates = entry_carbohydrates.get("1.0", tk.END).strip()
    dietary_fiber = entry_dietary_fiber.get("1.0", tk.END).strip()
    sugar = entry_sugar.get("1.0", tk.END).strip()
    protein = entry_protein.get("1.0", tk.END).strip()

    if not os.path.exists('./Documents/menu'):
        os.makedirs('./Documents/menu')
    
    doc = Document()
    
    doc.add_paragraph(f"{item_name} calories: {calories} calories, {calories_from_fat} calories from fat")
    doc.add_paragraph(f"{item_name} calories from fat: {calories_from_fat} calories from fat")
    doc.add_paragraph(f"{item_name} fat: {calories_from_fat} calories from fat, {fat} grams of total fat, {saturated_fat} grams of saturated fat, {trans_fat} grams of trans fat")
    doc.add_paragraph(f"{item_name} saturated fat: {saturated_fat} grams of saturated fat")
    doc.add_paragraph(f"{item_name} trans fat: {trans_fat} grams of trans fat")
    doc.add_paragraph(f"{item_name} total fat: {fat} grams of total fat")
    doc.add_paragraph(f"{item_name} cholesterol: {cholesterol} milligrams of cholesterol")
    doc.add_paragraph(f"{item_name} sodium: {sodium} milligrams of sodium")
    doc.add_paragraph(f"{item_name} carbohydrates: {carbohydrates} grams of carbohydrates")
    doc.add_paragraph(f"{item_name} dietary fiber: {dietary_fiber} grams of fiber")
    doc.add_paragraph(f"{item_name} sugar: {sugar} grams of sugar")
    doc.add_paragraph(f"{item_name} protein: {protein} grams of protein")
    
    file_path = f"./Documents/menu/{item_name}_nutrition_info.docx"
    
    doc.save(file_path)
    messagebox.showinfo("Success", f"nutrition_{item_name}.docx has been created successfully in './Documents/menu'!")

def focus_next_widget(event):
    event.widget.tk_focusNext().focus()
    return("break")

def focus_prev_widget(event):
    event.widget.tk_focusPrev().focus()
    return("break")

root = tk.Tk()
root.title("Nutritional Pal")
root.geometry("600x600")

pad_options = {'padx': 10, 'pady': 5}

tk.Label(root, text="Item Name").grid(row=0, column=0, **pad_options)
entry_item_name = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_item_name.grid(row=0, column=1, **pad_options)

tk.Label(root, text="Calories").grid(row=1, column=0, **pad_options)
entry_calories = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_calories.grid(row=1, column=1, **pad_options)

tk.Label(root, text="Calories from Fat").grid(row=2, column=0, **pad_options)
entry_calories_from_fat = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_calories_from_fat.grid(row=2, column=1, **pad_options)

tk.Label(root, text="Fat (grams)").grid(row=3, column=0, **pad_options)
entry_fat = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_fat.grid(row=3, column=1, **pad_options)

tk.Label(root, text="Saturated Fat (grams)").grid(row=4, column=0, **pad_options)
entry_saturated_fat = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_saturated_fat.grid(row=4, column=1, **pad_options)

tk.Label(root, text="Trans Fat (grams)").grid(row=5, column=0, **pad_options)
entry_trans_fat = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_trans_fat.grid(row=5, column=1, **pad_options)

tk.Label(root, text="Cholesterol (milligrams)").grid(row=7, column=0, **pad_options)
entry_cholesterol = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_cholesterol.grid(row=7, column=1, **pad_options)

tk.Label(root, text="Sodium (milligrams)").grid(row=8, column=0, **pad_options)
entry_sodium = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_sodium.grid(row=8, column=1, **pad_options)

tk.Label(root, text="Carbohydrates (grams)").grid(row=9, column=0, **pad_options)
entry_carbohydrates = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_carbohydrates.grid(row=9, column=1, **pad_options)

tk.Label(root, text="Dietary Fiber (grams)").grid(row=10, column=0, **pad_options)
entry_dietary_fiber = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_dietary_fiber.grid(row=10, column=1, **pad_options)

tk.Label(root, text="Sugar (grams)").grid(row=11, column=0, **pad_options)
entry_sugar = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_sugar.grid(row=11, column=1, **pad_options)

tk.Label(root, text="Protein (grams)").grid(row=12, column=0, **pad_options)
entry_protein = tk.Text(root, height=2, width=40, highlightthickness=0)
entry_protein.grid(row=12, column=1, **pad_options)

entries = [entry_item_name, entry_calories, entry_calories_from_fat, entry_fat, entry_saturated_fat, entry_trans_fat, entry_cholesterol, entry_sodium, entry_carbohydrates, entry_dietary_fiber, entry_sugar, entry_protein]

for entry in entries:
    entry.bind("<Down>", focus_next_widget)
    entry.bind("<Up>", focus_prev_widget)
    entry.bind("<Return>", save_to_docx)

save_button = tk.Button(root, text="Save", command=save_to_docx)
save_button.grid(row=14, column=0, columnspan=2, pady=10)

# Set focus to the item name entry
entry_item_name.focus()

root.mainloop()
